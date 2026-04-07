import os
import json
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

load_dotenv()

class ResumeBuilderService:
    def __init__(self):
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY is missing")
        self.client = OpenAI(api_key=api_key)

    def generate_targeted_updates(self, resume_text: str, jd_text: str) -> dict:
        """
        Tells the LLM to find 3-5 weak bullet points and rewrite them.
        Returns a mapping of { "original_exact_text": "new_star_text" }
        """
        prompt = f"""
        You are an elite executive resume writer.
        I will provide a candidate's resume and a target Job Description.
        
        Your task is to identify 3 to 5 specific bullet points in the resume that are weak or poorly aligned with the Job Description.
        Rewrite ONLY those specific bullet points to be highly impactful using the STAR method, injecting relevant keywords from the JD.
        
        Job Description: {jd_text[:1500]}
        
        Resume Text: {resume_text}

        CRITICAL INSTRUCTIONS:
        Return a strict JSON object where the keys are the EXACT original sentences from the resume (word-for-word, so I can do a string-replace), and the values are your new, improved versions.
        
        Example format:
        {{
            "Developed APIs for the backend system.": "Architected scalable RESTful APIs using FastAPI, improving data retrieval speeds by 40%.",
            "Worked on fixing bugs in the database.": "Spearheaded the resolution of critical SQL database anomalies, ensuring 99.9% uptime."
        }}
        """

        response = self.client.chat.completions.create(
            model="gpt-3.5-turbo", # Use gpt-4o for best results in production
            messages=[
                {"role": "system", "content": "You output strict JSON mappings for find-and-replace."},
                {"role": "user", "content": prompt}
            ],
            response_format={ "type": "json_object" },
            temperature=0.2
        )
        
        return json.loads(response.choices[0].message.content)

    def update_docx_resume(self, input_docx_path: str, output_docx_path: str, replacements: dict):
        """
        Opens the original DOCX, finds the old text, replaces it with the new text, 
        and saves it without breaking the user's template formatting.
        """
        doc = Document(input_docx_path)
        
        # Iterate through every paragraph in the document
        for para in doc.paragraphs:
            for old_text, new_text in replacements.items():
                # If the old text is found in this paragraph
                if old_text.strip() in para.text:
                    # We do a direct replace. 
                    # Note: python-docx can sometimes lose inline styling (like one bold word) 
                    # when replacing full text, but it preserves the paragraph's overall style/font.
                    para.text = para.text.replace(old_text.strip(), new_text.strip())
                    
        # Iterate through tables (a lot of resume templates use hidden tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text.strip() in para.text:
                                para.text = para.text.replace(old_text.strip(), new_text.strip())

        doc.save(output_docx_path)
        return output_docx_path

# Export singleton
resume_builder = ResumeBuilderService()